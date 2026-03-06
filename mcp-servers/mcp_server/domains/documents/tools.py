from mcp.server.fastmcp import FastMCP
from typing import Dict, Any, List
import os
import logging
from datetime import datetime

# File libraries
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib import colors
from reportlab.lib.units import inch

from docx import Document
from openpyxl import Workbook
from pathlib import Path

logger = logging.getLogger(__name__)

# Get user's Downloads folder
OUTPUT_DIR = Path.home() / "Downloads"

# Ensure folder exists (normally it does)
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def _generate_filename(prefix: str, extension: str) -> str:
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"{prefix}_{timestamp}.{extension}"


def register_document_tools(mcp: FastMCP):

    # =========================
    # Plain Text (.txt)
    # =========================
    @mcp.tool(
        description="""
        Create a plain text file with a .txt extension.

        Usage:
        - The user requested the output to be a plain text file or a document with txt extension.

        Parameters:
        - filename_prefix: Prefix of the file name.
        - content: Text content to write into the file.
        """
    )
    def create_txt_file(filename_prefix: str, content: str) -> Dict[str, Any]:
        filename = _generate_filename(filename_prefix, "txt")
        path = os.path.join(OUTPUT_DIR, filename)

        with open(path, "w", encoding="utf-8") as f:
            f.write(content)

        logger.info(f"TXT file created: {path}")
        return {
            "status": "success",
            "file_path": path,
            "file_name": filename,
            "file_type": "txt"
        }

    # # =========================
    # # Markdown (.md)
    # # =========================
    # @mcp.tool(
    #     description="""
    #     Create a Markdown (.md) file.

    #     Parameters:
    #     - filename_prefix: Prefix of the file name.
    #     - content: Markdown formatted content.
    #     """
    # )
    # def create_markdown_file(filename_prefix: str, content: str) -> Dict[str, Any]:
    #     filename = _generate_filename(filename_prefix, "md")
    #     path = os.path.join(OUTPUT_DIR, filename)

    #     with open(path, "w", encoding="utf-8") as f:
    #         f.write(content)

    #     logger.info(f"Markdown file created: {path}")
    #     return {
    #         "status": "success",
    #         "file_path": path,
    #         "file_name": filename,
    #         "file_type": "md"
    #     }

    # # =========================
    # # PDF (.pdf)
    # # =========================
    # @mcp.tool(
    #     description="""
    #     Create a simple PDF document.

    #     Parameters:
    #     - filename_prefix: Prefix of the file name.
    #     - title: Title of the document.
    #     - paragraphs: List of text paragraphs.
    #     """
    # )
    # def create_pdf_file(
    #     filename_prefix: str,
    #     title: str,
    #     paragraphs: List[str]
    # ) -> Dict[str, Any]:

    #     filename = _generate_filename(filename_prefix, "pdf")
    #     path = os.path.join(OUTPUT_DIR, filename)

    #     doc = SimpleDocTemplate(path)
    #     elements = []

    #     styles = getSampleStyleSheet()
    #     title_style = styles["Heading1"]
    #     body_style = styles["BodyText"]

    #     elements.append(Paragraph(title, title_style))
    #     elements.append(Spacer(1, 0.5 * inch))

    #     for p in paragraphs:
    #         elements.append(Paragraph(p, body_style))
    #         elements.append(Spacer(1, 0.2 * inch))

    #     doc.build(elements)

    #     logger.info(f"PDF file created: {path}")
    #     return {
    #         "status": "success",
    #         "file_path": path,
    #         "file_name": filename,
    #         "file_type": "pdf"
    #     }

    # =========================
    # Word (.docx)
    # =========================
    @mcp.tool(
        description="""
        Create a formated Word (.docx) document.

        Usage:
        - The user requested the output to be a Word document or a document with docx extension.

        Parameters:
        - filename_prefix: Prefix of the file name.
        - title: Document title.
        - paragraphs: List of paragraphs.
        """
    )
    def create_word_file(
        filename_prefix: str,
        title: str,
        paragraphs: List[str]
    ) -> Dict[str, Any]:

        filename = _generate_filename(filename_prefix, "docx")
        path = os.path.join(OUTPUT_DIR, filename)

        doc = Document()
        doc.add_heading(title, level=1)

        for p in paragraphs:
            doc.add_paragraph(p)

        doc.save(path)

        logger.info(f"Word file created: {path}")
        return {
            "status": "success",
            "file_path": path,
            "file_name": filename,
            "file_type": "docx"
        }

    # =========================
    # Excel (.xlsx)
    # =========================
    # @mcp.tool(
    #     description="""
    #     Create an Excel (.xlsx) file.

    #     Parameters:
    #     - filename_prefix: Prefix of the file name.
    #     - sheet_name: Name of the sheet.
    #     - headers: List of column headers.
    #     - rows: List of row values (each row is a list).
    #     """
    # )
    # def create_excel_file(
    #     filename_prefix: str,
    #     sheet_name: str,
    #     headers: List[str],
    #     rows: List[List[Any]]
    # ) -> Dict[str, Any]:

    #     filename = _generate_filename(filename_prefix, "xlsx")
    #     path = os.path.join(OUTPUT_DIR, filename)

    #     wb = Workbook()
    #     ws = wb.active
    #     ws.title = sheet_name

    #     ws.append(headers)
    #     for row in rows:
    #         ws.append(row)

    #     wb.save(path)

    #     logger.info(f"Excel file created: {path}")
    #     return {
    #         "status": "success",
    #         "file_path": path,
    #         "file_name": filename,
    #         "file_type": "xlsx"
    #     }